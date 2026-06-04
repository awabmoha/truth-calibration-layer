"""Build the TCL-v0 research writeup DOCX from the Markdown source.

This script intentionally keeps the paper source in Markdown and makes the
DOCX/PDF artifact reproducible.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "TCL-v0-research-writeup.md"
OUTPUT_DOCX = ROOT / "TCL-v0-research-writeup.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
TABLE_FILL = "F2F4F7"
BORDER = "A6B4C2"
CONTENT_WIDTH = Inches(6.5)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def apply_paragraph_spacing(paragraph, before=0, after=6, line=1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_formatted_text(paragraph, text: str, *, code=False) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=10, color=INK)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, name="Consolas" if code else "Calibri", size=10 if code else 11)


def add_body_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    apply_paragraph_spacing(paragraph)
    add_formatted_text(paragraph, text)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="TCL Bullet")
    add_formatted_text(paragraph, text)
    return paragraph


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="TCL Code")
    add_formatted_text(paragraph, "\n".join(lines), code=True)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        rows.append(cells)
    return rows


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = column_widths(col_count)
    for row_index, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            value = rows[row_index][col_index] if col_index < len(rows[row_index]) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_spacing(paragraph, after=0, line=1.0)
            run = paragraph.add_run(value)
            set_run_font(run, size=8 if col_count >= 6 else 9, bold=row_index == 0)
    doc.add_paragraph()


def column_widths(col_count: int) -> list[int]:
    if col_count == 7:
        return [1720, 1140, 1140, 1040, 1040, 1230, 1230]
    if col_count == 5:
        return [2040, 1740, 1740, 1420, 1420]
    width = 9360 // col_count
    widths = [width] * col_count
    widths[-1] += 9360 - sum(widths)
    return widths


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = doc.styles.add_style("TCL Bullet", 1)
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    code = doc.styles.add_style("TCL Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(10)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Truth Calibration Layer (TCL) - TCL-v0 Research Note")
    set_run_font(run, size=9, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Preliminary empirical note")
    set_run_font(run, size=9, color="667085")


def build_document() -> Document:
    doc = Document()
    setup_styles(doc)
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    seen_title = False

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, parse_table(table_lines))
            table_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(line)
            continue

        flush_table()

        if not stripped:
            continue

        if stripped.startswith("# "):
            title_text = stripped[2:].strip()
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_text(paragraph, title_text)
            seen_title = True
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
            continue

        paragraph = add_body_paragraph(doc, stripped)
        if seen_title and stripped.startswith("**") and ":**" in stripped:
            paragraph.paragraph_format.space_after = Pt(2)

    flush_table()
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
